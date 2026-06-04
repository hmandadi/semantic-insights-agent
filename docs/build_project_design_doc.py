from __future__ import annotations

from datetime import date
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSETS_DIR = DOCS_DIR / "assets"
OUTPUT_DOCX = DOCS_DIR / "Semantic_Insights_Agent_Project_Design_and_Implementation.docx"


COLORS = {
    "ink": "0B2545",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "muted": "667085",
    "light_fill": "F2F4F7",
    "pale_blue": "E8EEF5",
    "line": "D0D5DD",
    "green": "D9EAD3",
    "amber": "FFF4D6",
    "white": "FFFFFF",
}


def font(size: int, bold: bool = False):
    candidates = [
        "arialbd.ttf" if bold else "arial.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            pass
    return ImageFont.load_default()


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D0D5DD", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_fixed(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(row.cells[idx])
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_run(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, size=None, bold=None, color=None, italic=None, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_caption(doc, text):
    p = add_para(doc, text, size=9, italic=True, color=COLORS["muted"], after=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_hyperstyle_header(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Semantic Insights Agent | Project Design and Implementation")
    set_run(r, size=9, color=COLORS["muted"])
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Professional project design document")
    set_run(fr, size=9, color=COLORS["muted"])


def add_title_page(doc):
    section = doc.sections[0]
    add_hyperstyle_header(section)
    add_para(doc, "PROJECT DESIGN AND IMPLEMENTATION DOCUMENT", size=11, bold=True, color=COLORS["blue"], after=8)
    title = add_para(doc, "Semantic Insights Agent v1.0", size=28, bold=True, color=COLORS["ink"], after=6)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = add_para(
        doc,
        "Enterprise semantic analytics assistant for natural-language questions, governed SQL generation, PostgreSQL execution, charting, and business insight generation.",
        size=12,
        color="344054",
        after=18,
    )
    subtitle.paragraph_format.line_spacing = 1.2

    table = doc.add_table(rows=5, cols=2)
    set_table_fixed(table, [1.5, 4.75])
    data = [
        ("Prepared for", "Semantic Insights Agent project repository"),
        ("Prepared by", "Codex project documentation assistant"),
        ("Document date", date.today().strftime("%B %d, %Y")),
        ("Source reviewed", "app.py, src/*.py, config/semantic_manifest.yaml, README.md, test scripts"),
        ("Design preset", "Standard business brief"),
    ]
    for row, (label, value) in zip(table.rows, data):
        set_cell_shading(row.cells[0], COLORS["pale_blue"])
        row.cells[0].paragraphs[0].add_run(label)
        row.cells[1].paragraphs[0].add_run(value)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_run(run, size=10.5, bold=(cell == row.cells[0]), color=COLORS["ink"])
    add_para(doc, "", after=8)

    p = add_para(
        doc,
        "Executive takeaway: the project implements a clear MVP architecture: Streamlit receives the business question, AgentService creates the shared AgentState, SemanticService loads the YAML business model, LangGraph runs three nodes, and the UI renders SQL, results, charts, and insight. The state module is active as the typed state contract for the workflow even though it is not an executable node.",
        size=11,
        bold=True,
        color=COLORS["ink"],
        after=10,
    )
    p.paragraph_format.line_spacing = 1.18
    doc.add_page_break()


def add_manual_contents(doc):
    add_para(doc, "Contents", style="Heading 1")
    sections = [
        "1. Executive Summary",
        "2. Project Purpose and Scope",
        "3. Architecture Overview",
        "4. Runtime Request Flow",
        "5. LangGraph Workflow and AgentState Contract",
        "6. Component Design and File Responsibilities",
        "7. Semantic Layer and Data Model",
        "8. LLM, Prompting, and SQL Governance",
        "9. UI, Visualization, and Output Experience",
        "10. Configuration, Deployment, and Operations",
        "11. Testing and Validation",
        "12. Risks, Gaps, and Recommended Enhancements",
        "13. Implementation Appendix",
    ]
    for item in sections:
        add_para(doc, item, size=10.5, color=COLORS["ink"], after=2)
    doc.add_page_break()


def draw_box(draw, xy, title, body="", fill="#FFFFFF", outline="#2E74B5", title_fill=None):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=2)
    if title_fill:
        draw.rounded_rectangle((x1, y1, x2, y1 + 34), radius=14, fill=title_fill, outline=outline, width=0)
        draw.rectangle((x1, y1 + 20, x2, y1 + 34), fill=title_fill)
    title_font = font(20, True)
    body_font = font(15)
    draw.text((x1 + 16, y1 + 10), title, fill="#0B2545", font=title_font)
    if body:
        y = y1 + 45
        for paragraph in str(body).splitlines():
            wrapped = wrap(paragraph, width=max(18, int((x2 - x1) / 10))) or [""]
            for line in wrapped:
                draw.text((x1 + 16, y), line, fill="#344054", font=body_font)
                y += 20


def arrow(draw, start, end, color="#475467"):
    draw.line([start, end], fill=color, width=3)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - 12 * direction, ey - 7), (ex - 12 * direction, ey + 7)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 7, ey - 12 * direction), (ex + 7, ey - 12 * direction)]
    draw.polygon(pts, fill=color)


def path_arrow(draw, points, color="#475467"):
    draw.line(points, fill=color, width=3)
    arrow(draw, points[-2], points[-1], color=color)


def center_text(draw, xy, text, size=16, bold=False, fill="#0B2545"):
    fnt = font(size, bold)
    bbox = draw.textbbox((0, 0), text, font=fnt)
    x1, y1, x2, y2 = xy
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    draw.text((x1 + (x2 - x1 - tw) / 2, y1 + (y2 - y1 - th) / 2), text, font=fnt, fill=fill)


def create_architecture_diagram():
    path = ASSETS_DIR / "architecture_flow.png"
    img = Image.new("RGB", (1500, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((40, 32), "Figure 1. End-to-End Semantic Insights Agent Architecture", font=font(26, True), fill="#0B2545")

    boxes = {
        "user": (60, 120, 320, 220),
        "app": (410, 100, 720, 240),
        "agent": (820, 105, 1130, 235),
        "yaml": (1180, 90, 1450, 210),
        "graph": (420, 340, 760, 505),
        "state": (70, 370, 320, 505),
        "nodes": (850, 315, 1445, 545),
        "llm": (880, 635, 1120, 760),
        "db": (1205, 635, 1450, 760),
        "viz": (420, 635, 730, 760),
        "output": (70, 640, 320, 780),
    }
    draw_box(draw, boxes["user"], "User", "Business question", fill="#F9FAFB", title_fill="#E8EEF5")
    draw_box(draw, boxes["app"], "app.py", "Streamlit UI, example prompts, results table, KPI cards, charts, insight display.", fill="#FFFFFF", title_fill="#E8EEF5")
    draw_box(draw, boxes["agent"], "AgentService", "Creates initial AgentState and invokes the compiled LangGraph workflow.", fill="#FFFFFF", title_fill="#E8EEF5")
    draw_box(draw, boxes["yaml"], "semantic_manifest.yaml", "Tables, metrics, dimensions, and joins for Enterprise Sales Analytics.", fill="#FFF4D6", outline="#C08A00", title_fill="#FFF4D6")
    draw_box(draw, boxes["graph"], "graph.py", "StateGraph(AgentState): sql_generator -> sql_executor -> insight.", fill="#FFFFFF", title_fill="#E8EEF5")
    draw_box(draw, boxes["state"], "src/state.py", "TypedDict contract shared across AgentService, graph, and every node.", fill="#D9EAD3", outline="#5B8C5A", title_fill="#D9EAD3")
    draw_box(draw, boxes["nodes"], "nodes.py", "sql_generator_node uses semantic context and LLM. sql_executor_node calls DB. insight_node calls LLM.", fill="#FFFFFF", title_fill="#E8EEF5")
    draw_box(draw, boxes["llm"], "LLMService", "OpenAI-compatible API via OpenRouter/OpenAI base URL.", fill="#F9FAFB", title_fill="#E8EEF5")
    draw_box(draw, boxes["db"], "db.py", "Neon PostgreSQL connection, SELECT validation, pandas DataFrame output.", fill="#F9FAFB", title_fill="#E8EEF5")
    draw_box(draw, boxes["viz"], "VisualizationService", "Normalizes DataFrame and selects bar or line chart columns.", fill="#F9FAFB", title_fill="#E8EEF5")
    draw_box(draw, boxes["output"], "User Output", "Generated SQL, query results, visualization, business insight.", fill="#F9FAFB", title_fill="#E8EEF5")

    arrow(draw, (320, 170), (410, 170))
    arrow(draw, (720, 170), (820, 170))
    arrow(draw, (1130, 150), (1180, 150))
    arrow(draw, (975, 235), (610, 340))
    arrow(draw, (320, 435), (420, 435), color="#5B8C5A")
    arrow(draw, (760, 430), (850, 430))
    arrow(draw, (1035, 545), (1000, 635))
    arrow(draw, (1240, 545), (1325, 635))
    path_arrow(draw, [(720, 210), (785, 210), (785, 700), (730, 700)])
    arrow(draw, (420, 700), (320, 700))
    draw.text((78, 525), "State is not a separate runtime node; it is the schema passed through the flow.", font=font(17, True), fill="#326B35")
    img.save(path, quality=95)
    return path


def create_langgraph_diagram():
    path = ASSETS_DIR / "langgraph_state_flow.png"
    img = Image.new("RGB", (1500, 820), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((40, 32), "Figure 2. Current LangGraph Node Flow with AgentState Contract", font=font(26, True), fill="#0B2545")

    y = 170
    xs = [80, 330, 640, 950, 1260]
    labels = [
        ("START", "Entry point"),
        ("sql_generator_node", "question + semantic_context -> generated_sql"),
        ("sql_executor_node", "generated_sql -> query_result"),
        ("insight_node", "question + query_result -> answer"),
        ("END", "Final state returned"),
    ]
    for x, (title, body) in zip(xs, labels):
        draw_box(draw, (x, y, x + 190, y + 135), title, body, fill="#FFFFFF", title_fill="#E8EEF5")
    for idx in range(len(xs) - 1):
        arrow(draw, (xs[idx] + 190, y + 68), (xs[idx + 1], y + 68))

    draw.rounded_rectangle((120, 440, 1380, 700), radius=18, fill="#D9EAD3", outline="#5B8C5A", width=2)
    draw.text((150, 468), "AgentState from src/state.py", font=font(24, True), fill="#0B2545")
    state_items = [
        ("question", "Natural-language business question from Streamlit"),
        ("semantic_context", "YAML manifest loaded by SemanticService"),
        ("generated_sql", "SQL generated by LLMService"),
        ("query_result", "Rows returned by db.execute_query as list[dict]"),
        ("answer", "Business insight generated by LLMService"),
    ]
    start_x = 150
    start_y = 525
    col_width = 240
    for idx, (field, meaning) in enumerate(state_items):
        x1 = start_x + idx * col_width
        draw.rounded_rectangle((x1, start_y, x1 + 210, start_y + 125), radius=12, fill="#FFFFFF", outline="#8AB982", width=2)
        draw.text((x1 + 12, start_y + 12), field, font=font(19, True), fill="#0B2545")
        yy = start_y + 42
        for line in wrap(meaning, 24):
            draw.text((x1 + 12, yy), line, font=font(14), fill="#344054")
            yy += 18
    draw.text((150, 720), "Interpretation: AgentState is connected to the main flow through imports, graph schema declaration, initial state creation, and node input/output mutation.", font=font(18, True), fill="#326B35")
    img.save(path, quality=95)
    return path


def create_semantic_model_diagram():
    path = ASSETS_DIR / "semantic_data_model.png"
    img = Image.new("RGB", (1500, 860), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((40, 32), "Figure 3. Semantic Manifest and Sales Data Model", font=font(26, True), fill="#0B2545")

    fact = (560, 250, 940, 520)
    draw_box(draw, fact, "fact_sales", "sale_id PK\ncustomer_id FK\nproduct_id FK\nregion_id FK\nrevenue\nquantity", fill="#FFF4D6", outline="#C08A00", title_fill="#FFF4D6")
    dims = {
        "dim_customer": ((90, 170, 430, 365), "customer_id PK\ncustomer_name\ncustomer_segment"),
        "dim_product": ((90, 545, 430, 740), "product_id PK\nproduct_name\nproduct_category"),
        "dim_region": ((1060, 355, 1400, 550), "region_id PK\nregion_name\ncountry"),
    }
    for name, (xy, body) in dims.items():
        draw_box(draw, xy, name, body, fill="#FFFFFF", title_fill="#E8EEF5")
    arrow(draw, (560, 335), (430, 270))
    arrow(draw, (560, 450), (430, 645))
    arrow(draw, (940, 385), (1060, 450))

    draw.rounded_rectangle((560, 620, 940, 770), radius=16, fill="#D9EAD3", outline="#5B8C5A", width=2)
    draw.text((590, 646), "Business semantics", font=font(22, True), fill="#0B2545")
    for idx, text in enumerate(["Metrics: Revenue, Quantity Sold", "Dimensions: Customer, Product, Region", "Domain: Enterprise Sales Analytics"]):
        draw.text((590, 686 + idx * 24), text, font=font(16), fill="#344054")
    arrow(draw, (750, 620), (750, 520), color="#5B8C5A")
    img.save(path, quality=95)
    return path


def add_image(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.45))
    add_caption(doc, caption)


def add_table(doc, headers, rows, widths, header_fill=COLORS["light_fill"]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_fixed(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(header)
        set_run(r, size=9.5, bold=True, color=COLORS["ink"])
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(value)
            set_run(r, size=9.2, color="101828")
    add_para(doc, "", after=4)
    return table


def add_section_break(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_hyperstyle_header(doc.sections[-1])


def build_document():
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    architecture_path = create_architecture_diagram()
    langgraph_path = create_langgraph_diagram()
    semantic_path = create_semantic_model_diagram()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("101828")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    add_title_page(doc)
    add_manual_contents(doc)

    add_para(doc, "1. Executive Summary", style="Heading 1")
    add_para(
        doc,
        "Semantic Insights Agent v1.0 is an enterprise analytics MVP that converts natural-language business questions into governed PostgreSQL SELECT queries, executes those queries against a Neon PostgreSQL database, visualizes results in Streamlit, and uses an LLM to produce business-language insights.",
    )
    add_para(
        doc,
        "The implemented architecture is intentionally compact: a Streamlit presentation layer delegates workflow execution to AgentService; AgentService loads semantic metadata and creates AgentState; LangGraph coordinates SQL generation, SQL execution, and insight generation; VisualizationService handles chart preparation after the graph returns. The semantic manifest is the business governance layer that constrains table, metric, dimension, and join vocabulary.",
    )
    add_table(
        doc,
        ["Area", "Current implementation"],
        [
            ("Primary UI", "app.py Streamlit application with example buttons, text input, SQL display, DataFrame output, KPI cards, charting, and insight panel."),
            ("Workflow engine", "LangGraph StateGraph declared in src/graph.py using AgentState as the graph schema."),
            ("State contract", "src/state.py defines AgentState with question, semantic_context, generated_sql, query_result, and answer."),
            ("Semantic layer", "config/semantic_manifest.yaml models Enterprise Sales Analytics with fact_sales and customer/product/region dimensions."),
            ("LLM integration", "src/llm_service.py uses LangChain ChatOpenAI with OPENAI_API_KEY and optional OPENAI_BASE_URL for OpenRouter or another OpenAI-compatible endpoint."),
            ("Database", "src/db.py connects to PostgreSQL with SSL and validates that only SELECT statements are executed."),
        ],
        [1.75, 4.55],
    )

    add_para(doc, "2. Project Purpose and Scope", style="Heading 1")
    add_para(
        doc,
        "The project demonstrates how a semantic layer can mediate between business language and physical warehouse tables. Users ask questions such as 'Show revenue by region' or 'Show top 5 customers by revenue'. The system converts those questions into SQL using a governed manifest, executes the query, and presents both raw and interpreted outputs.",
    )
    add_para(doc, "In scope", style="Heading 2")
    add_table(
        doc,
        ["Capability", "Description"],
        [
            ("Natural-language analytics", "Business user asks a plain-language question in Streamlit."),
            ("Semantic SQL generation", "The LLM receives formatted table, metric, dimension, and relationship context before generating SQL."),
            ("Read-only query execution", "Generated SQL is validated as SELECT-only before execution."),
            ("Business insight generation", "Query rows are summarized into analyst-style prose."),
            ("Chart preparation", "Numeric, text, and datetime columns are inspected to select a chart type and x/y columns."),
        ],
        [1.8, 4.5],
    )
    add_para(doc, "Out of scope for v1.0", style="Heading 2")
    add_para(
        doc,
        "The current MVP does not yet include semantic manifest validation, SQL parser-based governance, user authentication, audit logging, query cost controls, streaming responses, conversational memory, formal API endpoints, or production deployment automation. These are appropriate next-phase enhancements.",
    )

    add_para(doc, "3. Architecture Overview", style="Heading 1")
    add_image(
        doc,
        architecture_path,
        "Figure 1 shows the complete architecture, including the corrected AgentState position as a shared contract rather than an executable node.",
    )
    add_para(
        doc,
        "The architecture follows a controller-orchestrator pattern. The UI does not directly call nodes, the database, or the LLM. Instead, it calls AgentService, which creates the initial state and invokes the compiled LangGraph. This keeps the Streamlit layer focused on presentation and allows the workflow implementation to evolve independently.",
    )

    add_para(doc, "4. Runtime Request Flow", style="Heading 1")
    add_table(
        doc,
        ["Step", "Module", "Runtime responsibility"],
        [
            ("1", "app.py", "Receives the user's business question through Streamlit."),
            ("2", "src/agent_service.py", "Instantiates AgentState with question, semantic_context, generated_sql, query_result, and answer."),
            ("3", "src/semantic_service.py", "Loads and exposes config/semantic_manifest.yaml as a dictionary."),
            ("4", "src/graph.py", "Invokes the compiled LangGraph workflow using StateGraph(AgentState)."),
            ("5", "src/nodes.py: sql_generator_node", "Formats semantic context and calls LLMService.generate_sql."),
            ("6", "src/nodes.py: sql_executor_node", "Calls db.execute_query and stores rows as list[dict]."),
            ("7", "src/nodes.py: insight_node", "Calls LLMService.generate_insight using the original question and query rows."),
            ("8", "src/agent_service.py", "Returns generated_sql, query_result, and answer to the UI."),
            ("9", "src/visualization_service.py", "Runs after workflow completion to prepare chart data and select chart type."),
            ("10", "app.py", "Displays generated SQL, table results, chart, and business insight."),
        ],
        [0.55, 1.8, 3.95],
    )

    add_para(doc, "5. LangGraph Workflow and AgentState Contract", style="Heading 1")
    add_image(
        doc,
        langgraph_path,
        "Figure 2 completes the current node flow and shows how src/state.py is connected through AgentState.",
    )
    add_para(
        doc,
        "src/state.py is connected to the main flow through three implementation points: AgentService imports AgentState and uses it for the initial_state type, graph.py creates StateGraph(AgentState), and nodes.py declares each node as accepting and returning AgentState. It is therefore not a missing flow node; it is the typed schema that LangGraph and the nodes share.",
    )
    add_table(
        doc,
        ["AgentState field", "Producer", "Consumer", "Purpose"],
        [
            ("question", "app.py -> AgentService", "sql_generator_node, insight_node", "Carries the original business question throughout the workflow."),
            ("semantic_context", "SemanticService.get_manifest", "sql_generator_node", "Supplies governed table, metric, dimension, and relationship metadata."),
            ("generated_sql", "sql_generator_node", "sql_executor_node, app.py", "Stores the LLM-generated PostgreSQL query."),
            ("query_result", "sql_executor_node", "insight_node, app.py, VisualizationService", "Stores database rows as records for display and insight generation."),
            ("answer", "insight_node", "app.py", "Stores final business-language summary."),
        ],
        [1.35, 1.55, 1.85, 1.65],
    )

    add_para(doc, "6. Component Design and File Responsibilities", style="Heading 1")
    add_table(
        doc,
        ["File", "Responsibility", "Key design notes"],
        [
            ("app.py", "Streamlit UI and presentation layer.", "Caches AgentService and VisualizationService, handles input validation, renders SQL/results/charts/insights, and catches exceptions."),
            ("src/agent_service.py", "Controller between UI and LangGraph.", "Builds initial AgentState, loads semantic manifest, invokes graph, returns selected result fields."),
            ("src/graph.py", "LangGraph workflow definition.", "Defines node ordering: sql_generator -> sql_executor -> insight -> END."),
            ("src/nodes.py", "Agentic execution steps.", "Formats manifest, calls LLM for SQL, executes SQL through db.py, calls LLM for insight."),
            ("src/state.py", "Shared state structure.", "Defines AgentState TypedDict and is the central contract across service, graph, and nodes."),
            ("src/semantic_service.py", "Semantic manifest access.", "Loads YAML once at service construction and provides helper accessors for metrics, dimensions, relationships, and full manifest."),
            ("src/prompts.py", "Prompt templates.", "Defines SQL generation rules and business insight instructions."),
            ("src/llm_service.py", "LLM provider client.", "Uses LangChain ChatOpenAI with environment-driven API key and base URL; strips SQL markdown fences."),
            ("src/db.py", "PostgreSQL execution layer.", "Loads DB environment variables, enforces SELECT-only validation, returns pandas DataFrame."),
            ("src/visualization_service.py", "Chart detection and data preparation.", "Normalizes numeric-looking strings, selects chart type, and identifies chart columns."),
            ("config/semantic_manifest.yaml", "Business semantic layer.", "Defines Enterprise Sales Analytics domain, tables, metrics, dimensions, and joins."),
        ],
        [1.55, 2.05, 2.7],
    )

    add_para(doc, "7. Semantic Layer and Data Model", style="Heading 1")
    add_image(
        doc,
        semantic_path,
        "Figure 3 summarizes the semantic manifest's fact/dimension sales model and business vocabulary.",
    )
    add_para(
        doc,
        "The semantic manifest is the bridge between business terminology and physical schema. It defines the fact table fact_sales, dimensions dim_customer, dim_product, and dim_region, and business metrics such as Revenue and Quantity Sold. SQL generation depends on this context to avoid invented table or column names.",
    )
    add_table(
        doc,
        ["Semantic element", "Implementation"],
        [
            ("Domain", "Enterprise Sales Analytics"),
            ("Fact table", "fact_sales with sale_id, customer_id, product_id, region_id, revenue, and quantity."),
            ("Dimensions", "customer, product, and region are mapped to dimension tables and business names."),
            ("Metrics", "revenue = SUM(fact_sales.revenue); quantity = SUM(fact_sales.quantity)."),
            ("Relationships", "fact_sales joins to dim_customer, dim_product, and dim_region using the corresponding foreign keys."),
        ],
        [1.7, 4.6],
    )

    add_para(doc, "8. LLM, Prompting, and SQL Governance", style="Heading 1")
    add_para(
        doc,
        "The SQL generation prompt instructs the model to generate a valid PostgreSQL SELECT query using only the semantic context. LLMService invokes ChatOpenAI with temperature 0.0, then strips markdown code fences if the model returns them. The insight prompt asks for business-language observations and explicitly keeps currency formatting in the UI layer.",
    )
    add_table(
        doc,
        ["Concern", "Current control", "Recommended next control"],
        [
            ("Table and column grounding", "Prompt rules plus semantic context formatting.", "Add deterministic semantic validation after SQL generation."),
            ("Read-only execution", "db.validate_sql requires SELECT and blocks major DML/DDL keywords.", "Replace keyword scanning with SQL parser validation and statement normalization."),
            ("Query safety", "Database credentials loaded from .env and SSL required.", "Add result limits, query timeout, role-based DB user, and audit logging."),
            ("LLM output cleanup", "_strip_markdown removes common fenced code blocks.", "Add retry/fix loop when SQL fails validation or execution."),
            ("Business insight", "Prompt separates result numeric formatting from UI formatting.", "Add structured insight schema if downstream automation is needed."),
        ],
        [1.65, 2.3, 2.35],
    )

    add_para(doc, "9. UI, Visualization, and Output Experience", style="Heading 1")
    add_para(
        doc,
        "The Streamlit UI gives users fast entry points through example buttons and a text input. After analysis, it presents generated SQL, summary metrics, tabular query results, optional charting, and a business insight. VisualizationService operates after the LangGraph flow and is intentionally separate from AgentState in the current implementation.",
    )
    add_table(
        doc,
        ["Output surface", "Behavior"],
        [
            ("Generated SQL", "Displayed in a SQL code block for transparency and debugging."),
            ("Summary metrics", "Rows returned, columns returned, and question length are shown as Streamlit metric cards."),
            ("Query results", "Converted to a pandas DataFrame and displayed with revenue number formatting when a revenue-like column is detected."),
            ("Visualization", "Bar chart for text plus numeric columns; line chart for datetime plus numeric columns; no chart for unsupported shapes."),
            ("Business insight", "LLM-generated narrative summary rendered as markdown."),
        ],
        [1.6, 4.7],
    )

    add_para(doc, "10. Configuration, Deployment, and Operations", style="Heading 1")
    add_table(
        doc,
        ["Configuration item", "Purpose"],
        [
            ("OPENAI_API_KEY", "Required by LLMService for OpenAI-compatible model access."),
            ("OPENAI_BASE_URL", "Optional base URL that enables OpenRouter or another compatible endpoint."),
            ("DB_HOST, DB_PORT, DB_NAME, DB_USER, DB_PASSWORD", "Required database connectivity variables loaded from .env."),
            ("config/semantic_manifest.yaml", "Versioned semantic layer file that governs analytics vocabulary."),
            ("requirements.txt", "Runtime dependencies: streamlit, langgraph, langchain, langchain-openai, psycopg2-binary, python-dotenv, pyyaml, pandas."),
        ],
        [2.05, 4.25],
    )
    add_para(
        doc,
        "A production deployment should use managed secrets rather than a local .env file, a least-privilege read-only database role, network restrictions to the database, centralized logs, and health checks for the LLM and database dependencies.",
    )

    add_para(doc, "11. Testing and Validation", style="Heading 1")
    add_para(
        doc,
        "The repository contains lightweight script-style tests for semantic loading, node execution, graph invocation, database execution, and end-to-end business insight generation. These are useful during MVP development but should be converted into automated pytest tests with mocks for LLM and database dependencies.",
    )
    add_table(
        doc,
        ["Test file", "Observed purpose", "Notes"],
        [
            ("test_semantic.py", "Instantiates SemanticService and prints a metric.", "Good smoke test for manifest loading."),
            ("test_node.py", "Calls sql_generator_node.", "Current script passes semantic_context from load_manifest(), which returns None; use get_manifest()."),
            ("test_graph.py", "Invokes the compiled graph with a sample question.", "Exercises all nodes and external dependencies."),
            ("test_db.py", "Executes a simple SELECT against dim_region.", "Requires live database credentials."),
            ("test_business_insight.py", "Runs generator, executor, and insight nodes sequentially.", "Useful end-to-end smoke script; should become an integration test."),
            ("text_to_sql.py", "Generates SQL and executes it.", "Development utility for validating natural-language to SQL behavior."),
        ],
        [1.45, 2.15, 2.7],
    )

    add_para(doc, "12. Risks, Gaps, and Recommended Enhancements", style="Heading 1")
    add_table(
        doc,
        ["Priority", "Gap", "Recommendation"],
        [
            ("High", "SQL governance relies on prompt adherence plus keyword validation.", "Add parser-based validation, allowlisted tables/columns, max row limits, and query timeout."),
            ("High", "LLM/database failures surface mainly as generic app exceptions.", "Add node-level error fields in AgentState and user-safe remediation messages."),
            ("Medium", "LLMService is instantiated inside each LLM node.", "Inject or cache the LLM service to reduce setup overhead and improve testability."),
            ("Medium", "Semantic manifest has no schema validation.", "Introduce a Pydantic or JSON Schema validator and CI check."),
            ("Medium", "Tests are script-based and call live dependencies.", "Add pytest, fixtures, mocked LLM responses, and mocked DataFrame query outputs."),
            ("Low", "Visualization is post-workflow only.", "Keep as-is for UI responsibility, or add a chart-spec field to AgentState if API clients need visualization metadata."),
        ],
        [0.75, 2.35, 3.2],
    )

    add_para(doc, "13. Implementation Appendix", style="Heading 1")
    add_para(doc, "Current LangGraph node flow", style="Heading 2")
    add_para(doc, "START -> sql_generator_node -> sql_executor_node -> insight_node -> END", size=11, bold=True, color=COLORS["ink"])
    add_para(doc, "How nodes.py works", style="Heading 2")
    add_table(
        doc,
        ["Node", "Input fields", "Uses", "Output field"],
        [
            ("sql_generator_node", "question, semantic_context", "Formatted semantic manifest, SQL_GENERATION_PROMPT, LLMService", "generated_sql"),
            ("sql_executor_node", "generated_sql", "db.execute_query", "query_result"),
            ("insight_node", "question, query_result", "INSIGHT_PROMPT, LLMService", "answer"),
        ],
        [1.5, 1.45, 2.35, 1.0],
    )
    add_para(doc, "Complete corrected architecture statement", style="Heading 2")
    add_para(
        doc,
        "src/state.py is connected to the project as the shared state schema. The flow starts in app.py, moves into AgentService, loads the semantic manifest through SemanticService, invokes the LangGraph compiled from graph.py, passes AgentState through nodes.py, calls LLMService and db.py, returns the final state to AgentService, and then app.py uses VisualizationService to prepare chart output.",
    )

    doc.core_properties.title = "Semantic Insights Agent Project Design and Implementation"
    doc.core_properties.subject = "Architecture and implementation documentation"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "Semantic Insights Agent, LangGraph, Streamlit, semantic layer, PostgreSQL, LLM"
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    out = build_document()
    print(out)
